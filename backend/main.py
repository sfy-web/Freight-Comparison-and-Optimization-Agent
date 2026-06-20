from fastapi import FastAPI, HTTPException, Request, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional, List
import os
import io
import logging
import pandas as pd

logger = logging.getLogger(__name__)

import config
from models import OrderRequest, ComparisonResult


class DocxExportRequest(OrderRequest):
    feedback: Optional[str] = None
    use_ai: Optional[bool] = True  # 是否使用了AI助手


# 数据源实例（启动时初始化）
data_store = None
freight_service = None
active_data_source = "unknown"


def _init_data_source():
    """根据配置初始化数据源，失败时自动降级到 CSV"""
    global data_store, freight_service, active_data_source

    if config.DATA_STORE == "sqlite":
        try:
            from database import init_db_if_needed, check_db_integrity, get_session_factory
            from freight_service import DBDataStore, FreightService

            # 尝试初始化数据库
            if not init_db_if_needed():
                raise RuntimeError("数据库初始化失败")

            if not check_db_integrity(config.EXPECTED_MIN_ROW_COUNT):
                raise RuntimeError("数据库完整性检查失败")

            session_factory = get_session_factory()
            data_store = DBDataStore(session_factory, auto_init=False)
            freight_service = FreightService(data_store)
            active_data_source = "sqlite"
            logger.info("数据源: SQLite")
            return
        except Exception as e:
            logger.error(f"SQLite 初始化失败，降级到 CSV: {e}")

    # CSV 模式或降级
    try:
        from freight_service import CSVDataStore, FreightService

        data_dir = config.CSV_DATA_DIR
        candidates = [
            os.path.join(data_dir, "FreightRates_combined.csv"),
            os.path.join(data_dir, "FreightRates_with_rating.csv"),
            os.path.join(data_dir, "FreightRates.csv"),
        ]
        csv_path = None
        for p in candidates:
            if os.path.exists(p):
                csv_path = p
                break

        if csv_path is None:
            raise FileNotFoundError(f"未找到 CSV 文件: {data_dir}")

        extended_path = os.path.join(data_dir, "FreightRates_extended.csv")
        data_store = CSVDataStore(csv_path, extended_path if csv_path != extended_path else None)
        freight_service = FreightService(data_store)
        active_data_source = "csv"
        logger.info(f"数据源: CSV ({csv_path})")
    except Exception as e:
        logger.error(f"CSV 初始化也失败: {e}")
        raise


# 初始化数据源
_init_data_source()

# 工具管理器导入
tool_manager = None
try:
    from tools import setup_tools, ToolManager
    tool_manager = setup_tools()
    print(f"工具管理器初始化成功，已注册 {len(tool_manager.tools)} 个工具")
except Exception as e:
    print(f"工具管理器初始化失败: {e}")

# LLM服务可选导入
llm_service = None
try:
    from llm_service import LLMService
    llm_service = LLMService(tool_manager=tool_manager)
except Exception as e:
    print(f"LLM服务初始化失败: {e}")

app = FastAPI(
    title="运输方案比价与优化智能体",
    description="根据订单信息自动匹配承运商报价，完成运费核算和多方案比价",
    version="2.0.0"
)

# CORS配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class ChatRequest(BaseModel):
    message: str
    system_prompt: Optional[str] = None
    session_id: Optional[str] = None


class ParseRequest(BaseModel):
    text: str
    session_id: Optional[str] = None


class ContinueRequest(BaseModel):
    session_id: str
    message: str


class TestLLMRequest(BaseModel):
    api_key: str
    base_url: str
    model: Optional[str] = "mimo-v2.5-pro"


def get_user_llm(request) -> object:
    """从请求 headers 中读取用户凭证，返回对应的 LLMService 实例。
    无用户凭证时返回全局 llm_service（本地 .env 配置）。"""
    user_key = request.headers.get("X-API-Key")
    user_url = request.headers.get("X-Base-URL")
    user_model = request.headers.get("X-Model")

    if user_key and user_url:
        from llm_service import LLMService
        return LLMService(
            tool_manager=tool_manager,
            api_key=user_key,
            base_url=user_url,
            model=user_model or "mimo-v2.5-pro",
        )
    return llm_service


# 挂载前端静态文件
STATIC_PATH = os.path.join(os.path.dirname(__file__), "static")
FRONTEND_PATH = os.path.join(os.path.dirname(__file__), "..", "frontend")

if os.path.exists(STATIC_PATH):
    app.mount("/assets", StaticFiles(directory=os.path.join(STATIC_PATH, "assets")), name="assets")
    SERVE_PATH = STATIC_PATH
elif os.path.exists(FRONTEND_PATH):
    SERVE_PATH = FRONTEND_PATH
else:
    SERVE_PATH = None


@app.get("/")
async def root():
    """返回前端页面"""
    if SERVE_PATH:
        index_path = os.path.join(SERVE_PATH, "index.html")
        if os.path.exists(index_path):
            return FileResponse(index_path)
    return {"message": "运输方案比价与优化智能体 API"}


@app.get("/api/ports")
async def get_ports():
    """获取可用港口列表"""
    return freight_service.get_available_ports()


@app.get("/api/statistics")
async def get_statistics():
    """获取数据统计信息"""
    stats = freight_service.get_statistics()
    stats["data_source"] = active_data_source
    return stats


@app.get("/api/cache_stats")
async def get_cache_stats():
    """获取缓存统计信息"""
    return freight_service.get_cache_stats()


REQUIRED_COLUMNS = [
    'Carrier', 'Orig_Port', 'Dest_Port', 'Min_Weight_Quant', 'Max_Weight_Quant',
    'Service_Level', 'Min_Cost', 'Rate', 'Mode_DSC', 'TPT_Day_Count', 'Carrier_Type'
]


@app.post("/api/upload_data")
async def upload_data(file: UploadFile = File(...)):
    """上传 CSV/Excel 文件替换数据库中的运价数据"""
    global data_store, freight_service

    # 校验文件类型
    filename = file.filename.lower()
    if not (filename.endswith('.csv') or filename.endswith('.xlsx') or filename.endswith('.xls')):
        raise HTTPException(status_code=400, detail="仅支持 CSV 和 Excel (.xlsx/.xls) 文件")

    try:
        content = await file.read()

        # 读取为 DataFrame
        if filename.endswith('.csv'):
            df = pd.read_csv(io.BytesIO(content))
        else:
            df = pd.read_excel(io.BytesIO(content))

        # 清洗
        from common.data_cleaner import clean_dataframe
        df = clean_dataframe(df)

        # 校验必需列
        missing_cols = [c for c in REQUIRED_COLUMNS if c not in df.columns]
        if missing_cols:
            raise HTTPException(
                status_code=400,
                detail=f"缺少必需列: {', '.join(missing_cols)}"
            )

        if len(df) == 0:
            raise HTTPException(status_code=400, detail="文件中没有有效数据")

        # 检测是否有服务评分
        has_rating = False
        if 'Service_Rating' in df.columns:
            unique_vals = set(df['Service_Rating'].dropna().unique()) - {'C'}
            has_rating = len(unique_vals) > 0

        # 导入数据库
        from database import get_session_factory
        from db_models import FreightRate

        Session = get_session_factory()
        session = Session()
        try:
            session.query(FreightRate).delete()
            session.commit()

            records = []
            for _, row in df.iterrows():
                records.append({
                    "carrier": row['Carrier'],
                    "orig_port": row['Orig_Port'],
                    "dest_port": row['Dest_Port'],
                    "min_weight": float(row['Min_Weight_Quant']),
                    "max_weight": float(row['Max_Weight_Quant']),
                    "service_level": row['Service_Level'],
                    "min_cost": float(row['Min_Cost']),
                    "rate": float(row['Rate']),
                    "mode": row['Mode_DSC'],
                    "transport_days": int(row['TPT_Day_Count']),
                    "carrier_type": row['Carrier_Type'],
                    "service_rating": row.get('Service_Rating', 'C'),
                })

            batch_size = config.IMPORT_BATCH_SIZE
            for i in range(0, len(records), batch_size):
                batch = records[i:i + batch_size]
                session.bulk_insert_mappings(FreightRate, batch)
                session.commit()

        except Exception as e:
            session.rollback()
            raise HTTPException(status_code=500, detail=f"数据导入失败: {str(e)}")
        finally:
            session.close()

        # 刷新 freight_service 的缓存和评分检测
        if hasattr(freight_service.data_store, 'refresh_service_rating'):
            freight_service.data_store.refresh_service_rating()
        freight_service.clear_compare_cache()

        # 重新创建数据源实例以刷新内存状态
        from database import get_session_factory
        from freight_service import DBDataStore, FreightService
        session_factory = get_session_factory()
        data_store = DBDataStore(session_factory, auto_init=False)
        data_store._detect_service_rating()
        freight_service = FreightService(data_store)

        carriers = sorted(df['Carrier'].unique().tolist())

        return {
            "success": True,
            "message": f"成功导入 {len(df)} 条记录",
            "total_records": len(df),
            "carriers": carriers,
            "has_service_rating": has_rating,
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件处理失败: {str(e)}")


@app.get("/api/data_preview")
async def data_preview(page: int = 1, page_size: int = 50):
    """分页浏览数据库中的运价数据（只读）"""
    from database import get_engine
    from sqlalchemy import text

    engine = get_engine()
    try:
        with engine.connect() as conn:
            # 总行数
            total = conn.execute(text("SELECT COUNT(*) FROM freight_rates")).scalar()

            # 列名
            columns = [
                'id', 'carrier', 'orig_port', 'dest_port',
                'min_weight', 'max_weight', 'service_level',
                'min_cost', 'rate', 'mode', 'transport_days',
                'carrier_type', 'service_rating'
            ]

            # 分页查询
            offset = (page - 1) * page_size
            rows = conn.execute(
                text(f"SELECT * FROM freight_rates LIMIT :limit OFFSET :offset"),
                {"limit": page_size, "offset": offset}
            ).fetchall()

            row_list = []
            for row in rows:
                row_dict = {}
                for i, col in enumerate(columns):
                    val = row[i]
                    if hasattr(val, '__float__'):
                        val = float(val)
                    row_dict[col] = val
                row_list.append(row_dict)

        return {
            "total": total,
            "page": page,
            "page_size": page_size,
            "columns": columns,
            "rows": row_list,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"数据查询失败: {str(e)}")


@app.post("/api/compare", response_model=ComparisonResult)
async def compare_freight(order: OrderRequest):
    """执行运费比价"""
    try:
        result = freight_service.compare(order)
        return result
    except Exception as e:
        error_msg = str(e)
        if "greater_than" in error_msg:
            raise HTTPException(status_code=400, detail="参数错误：货物重量必须大于0，请检查输入")
        elif "valid_ports" in error_msg or "PORT" in error_msg:
            raise HTTPException(status_code=400, detail="参数错误：港口代码无效，请使用有效的PORT代码")
        else:
            raise HTTPException(status_code=500, detail=f"系统错误：{error_msg}")


@app.post("/api/chat")
async def chat(request: ChatRequest, req: Request):
    """调用 LLM 进行对话"""
    active_llm = get_user_llm(req)
    if active_llm is None:
        return {"response": "LLM服务未配置", "model": "none", "configured": False}
    response = active_llm.chat(request.message, request.system_prompt)
    return {"response": response, "model": active_llm.model, "configured": active_llm.is_configured()}


@app.post("/api/test_llm")
async def test_llm(request: TestLLMRequest):
    """测试用户提供的 LLM 凭证是否可用"""
    try:
        from llm_service import LLMService
        test_service = LLMService(
            tool_manager=None,
            api_key=request.api_key,
            base_url=request.base_url,
            model=request.model or "mimo-v2.5-pro",
        )
        result = test_service.test_connection()
        return result
    except Exception as e:
        return {"success": False, "message": f"测试失败: {str(e)}", "model": ""}


@app.post("/api/agentic_chat")
async def agentic_chat(request: ChatRequest, req: Request):
    """
    Agentic 对话接口 v3 — LLM 主导解析 + 多轮槽位补全

    主路径: parse_agent_intent (LLM) → 意图路由 → 工具调用 / 推荐 / 追问
    Fallback: classify_intent (静态正则) 在 LLM 不可用/失败时使用

    返回统一结构:
    {
      "reply_type": "clarification | recommendation | no_result | error | general",
      "message": "面向用户的自然语言回复",
      "intent": "...",
      "missing_fields": [...],
      "order": {...},
      "recommendation": {...} | null,
      "plans": [...],
      "next_actions": [...],
      "tool_calls": [...], "tool_results": [...],    // 向后兼容
      "response": "...",                              // 向后兼容
      "session_id": "...",
      "parse_source": "llm | fallback",
    }
    """
    from models import (
        AgenticOrderInfo, AgenticChatResponse,
        RecommendationSummary, PlanSummary,
        OrderRequest
    )

    # 获取当前活跃的 LLM 服务（用户凭证优先，fallback 到本地 .env）
    active_llm = get_user_llm(req)

    # ---- 兜底: LLM 服务对象未初始化 ----
    if active_llm is None:
        return {
            "reply_type": "error",
            "message": "LLM服务未配置，请在 .env 文件中填写 DASHSCOPE_API_KEY",
            "intent": "general",
            "missing_fields": [], "order": None,
            "recommendation": None, "plans": [],
            "next_actions": ["配置 API Key 后重试"],
            "response": "LLM服务未配置",
            "tool_calls": [], "tool_results": [],
            "model": "none", "configured": False,
            "session_id": None, "parse_source": "fallback"
        }

    try:
        # ================================================================
        # 阶段 0: 多轮上下文恢复
        # ================================================================
        partial_order = None
        session = None
        if request.session_id and request.session_id in active_llm.sessions:
            session = active_llm.sessions[request.session_id]
            partial_order = session.partial_data if session.partial_data else None

        # ================================================================
        # 阶段 1: 意图解析
        # 优化：多轮补充字段时，直接用代码解析，不调用 LLM
        # ================================================================
        is_followup = (
            partial_order is not None
            and session is not None
            and session.missing_fields
            and session.current_state == "parsing"
        )

        if is_followup:
            # 多轮补充字段：用代码直接解析，避免 LLM 错误解析
            import re

            # 提取用户补充的字段
            followup_order = active_llm._extract_followup_slots(request.message, partial_order)

            # 合并到 partial_order：只更新用户新提供的字段
            order = dict(partial_order)
            for field in ("weight", "orig_port", "dest_port", "max_days", "priority"):
                if followup_order.get(field) is not None:
                    order[field] = followup_order[field]

            # 计算缺失字段
            missing = []
            if order.get("weight") is None:
                missing.append("weight")
            if order.get("orig_port") is None:
                missing.append("orig_port")
            if order.get("dest_port") is None:
                missing.append("dest_port")

            intent = "compare_freight"
            parse_source = "code_followup"
            llm_message = ""
        else:
            # 首次输入或无会话上下文：调用 LLM 解析
            classification = active_llm.parse_agent_intent(
                message=request.message,
                session_id=request.session_id,
                partial_order=partial_order,
            )

            intent = classification["intent"]
            order = classification["order"]
            missing = classification.get("missing_fields", [])
            parse_source = classification.get("parse_source", "fallback")
            llm_message = classification.get("message", "")

        # 构建 order 对象
        order_info = AgenticOrderInfo(
            weight=order.get("weight"),
            orig_port=order.get("orig_port"),
            dest_port=order.get("dest_port"),
            max_days=order.get("max_days"),
            priority=order.get("priority"),
        )

        # ================================================================
        # 阶段 2: 多轮会话管理
        # ================================================================
        new_session_id = request.session_id

        def _save_session():
            """将当前解析状态保存到会话（用于多轮追问）"""
            nonlocal new_session_id, session
            if new_session_id and new_session_id in active_llm.sessions:
                session = active_llm.sessions[new_session_id]
            else:
                import uuid
                new_session_id = str(uuid.uuid4())
                from llm_service import ConversationSession
                session = ConversationSession(session_id=new_session_id)
                active_llm.sessions[new_session_id] = session

            session.current_state = "parsing"
            session.partial_data = {
                "weight": order.get("weight"),
                "orig_port": order.get("orig_port"),
                "dest_port": order.get("dest_port"),
                "max_days": order.get("max_days"),
                "priority": order.get("priority"),
            }
            session.missing_fields = list(missing)
            session.user_feedback = llm_message

        def _clear_session():
            """信息完整后清理会话"""
            nonlocal new_session_id
            if new_session_id and new_session_id in active_llm.sessions:
                active_llm.sessions[new_session_id].current_state = "completed"

        # ================================================================
        # 阶段 3: 根据意图和缺失字段路由
        # ================================================================

        # ---- 3a: compare_freight 但缺字段 → clarification ----
        if intent == "compare_freight" and missing:
            _save_session()

            field_cn = {"weight": "货物重量", "orig_port": "起运港", "dest_port": "目的港"}
            missing_cn = [field_cn.get(f, f) for f in missing]
            known_parts = []
            if order["weight"]: known_parts.append(f"{order['weight']}kg")
            if order["orig_port"]: known_parts.append(f"起运港 {order['orig_port']}")
            if order["dest_port"]: known_parts.append(f"目的港 {order['dest_port']}")
            if order["max_days"]: known_parts.append(f"{order['max_days']}天内")

            if known_parts:
                message = f"已收到您的运输需求：{'，'.join(known_parts)}。\n\n还需要补充：{'、'.join(missing_cn)}。\n\n例如：从上海运100kg到深圳，5天内到达。"
            else:
                message = f"您好！请提供运输需求信息：{'、'.join(missing_cn)}。\n\n例如：从上海运100kg到深圳，5天内到达。"

            if llm_message and llm_message != message:
                message = llm_message  # 优先使用 LLM 生成的自然语言消息

            return {
                "reply_type": "clarification",
                "message": message,
                "intent": intent,
                "missing_fields": missing,
                "order": order_info.model_dump() if order_info else None,
                "recommendation": None, "plans": [],
                "next_actions": [f"请提供{'、'.join(missing_cn)}"],
                "response": message,
                "tool_calls": [], "tool_results": [],
                "model": active_llm.model, "configured": True,
                "session_id": new_session_id, "parse_source": parse_source
            }

        # ---- 3b: compare_freight 信息完整 → 执行比价 ----
        if intent == "compare_freight" and not missing:
            _clear_session()

            try:
                freight_order = OrderRequest(
                    weight=order["weight"],
                    orig_port=order["orig_port"],
                    dest_port=order["dest_port"],
                    max_days=order.get("max_days"),
                    priority=order.get("priority")
                )
            except Exception:
                return {
                    "reply_type": "error",
                    "message": f"订单参数无效：重量必须大于0，港口代码需在PORT02-PORT11范围内",
                    "intent": intent,
                    "missing_fields": missing,
                    "order": order_info.model_dump() if order_info else None,
                    "recommendation": None, "plans": [],
                    "next_actions": ["检查港口代码是否正确", "确认重量大于0"],
                    "response": "订单参数无效",
                    "tool_calls": [], "tool_results": [],
                    "model": active_llm.model, "configured": True,
                    "session_id": new_session_id, "parse_source": parse_source
                }

            result = freight_service.compare(freight_order)

            # 构建 plans 列表
            plans_data = []
            for plan in result.available_plans[:5]:
                plans_data.append(PlanSummary(
                    carrier=plan.carrier,
                    mode="空运" if plan.mode == "AIR" else "陆运",
                    service_level="门到门" if plan.service_level == "DTD" else "门到港",
                    transport_days=plan.transport_days,
                    total_cost=plan.total_cost,
                    service_rating=plan.service_rating,
                    score=plan.score
                ).model_dump())

            # ---- 3b-i: 有推荐方案 → recommendation ----
            if result.recommended_plan or result.transfer_routes:
                rp = result.recommended_plan.plan if result.recommended_plan else None
                rec = None
                if rp:
                    rec = RecommendationSummary(
                        carrier=rp.carrier,
                        transport_days=rp.transport_days,
                        total_cost=rp.total_cost,
                        mode="空运" if rp.mode == "AIR" else "陆运",
                        service_level="门到门" if rp.service_level == "DTD" else "门到港",
                        service_rating=rp.service_rating,
                        score=rp.score,
                        reason=result.recommended_plan.reason
                    )

                # 构建转运路线数据
                transfer_data = None
                if result.transfer_routes:
                    transfer_data = [t.model_dump() for t in result.transfer_routes]
                fallback_data = result.fallback_transfer.model_dump() if result.fallback_transfer else None

                # ---- LLM 反馈生成（阶段2） ----
                fb = active_llm.generate_agent_feedback(
                    user_message=request.message,
                    order=order,
                    recommendation=rec.model_dump() if rec else None,
                    plans=plans_data,
                    total_plans_found=result.total_plans_found,
                    scoring_weights=result.scoring_weights,
                    reply_type="recommendation" if rec else "general",
                    intent=intent,
                    parse_source=parse_source,
                    next_actions=[],
                    transfer_routes=transfer_data,
                    fallback_transfer=fallback_data,
                )
                message = fb["message"]
                feedback_source = fb["feedback_source"]
                feedback_reason = fb["feedback_reason"]

                next_actions = [
                    "查看全部方案详情",
                    "导出比价报告",
                    "调整时效要求重新查询" if not order.get("max_days") else "放宽时效要求重新查询"
                ]

                return {
                    "reply_type": "recommendation" if rec else "general",
                    "message": message,
                    "feedback_source": feedback_source,
                    "feedback_reason": feedback_reason,
                    "intent": intent,
                    "missing_fields": [],
                    "order": order_info.model_dump() if order_info else None,
                    "recommendation": rec.model_dump() if rec else None,
                    "plans": plans_data,
                    "transfer_routes": transfer_data,
                    "fallback_transfer": fallback_data,
                    "fallback_reason": result.fallback_reason or "",
                    "next_actions": next_actions,
                    "response": message,
                    "tool_calls": [{"tool": "compare_freight", "parameters": {
                        "weight": order["weight"], "orig_port": order["orig_port"],
                        "dest_port": order["dest_port"],
                        "max_days": order.get("max_days"), "priority": order.get("priority")
                    }}],
                    "tool_results": [{"success": True, "tool": "compare_freight",
                                    "total_plans": result.total_plans_found}],
                    "model": active_llm.model, "configured": True,
                    "session_id": new_session_id, "parse_source": parse_source
                }

            # ---- 3b-ii: 无方案 → no_result ----
            else:
                route_str = f"{order['orig_port']} -> {order['dest_port']}"
                has_time_filter = order.get("max_days") is not None

                if has_time_filter:
                    no_result_reason = f"在 {order['max_days']} 天时效内，{route_str} 路线无可用方案"
                    next_actions = [
                        f"尝试放宽时效要求至 {order['max_days'] + 3} 天",
                        "选择空运方式",
                        "查询其他港口组合"
                    ]
                else:
                    no_result_reason = f"{route_str} 路线暂无可用方案"
                    next_actions = [
                        "尝试其他目的港",
                        "调整货物重量",
                        "联系客服获取更多方案"
                    ]

                # ---- LLM 反馈生成（阶段2） ----
                fb = active_llm.generate_agent_feedback(
                    user_message=request.message,
                    order=order,
                    recommendation=None,
                    plans=[],
                    total_plans_found=0,
                    scoring_weights=None,
                    reply_type="no_result",
                    intent=intent,
                    parse_source=parse_source,
                    no_result_reason=no_result_reason,
                    next_actions=next_actions,
                )

                return {
                    "reply_type": "no_result",
                    "message": fb["message"],
                    "intent": intent,
                    "missing_fields": [],
                    "order": order_info.model_dump() if order_info else None,
                    "recommendation": None, "plans": [],
                    "next_actions": next_actions,
                    "response": fb["message"],
                    "tool_calls": [{"tool": "compare_freight", "parameters": {
                        "weight": order["weight"], "orig_port": order["orig_port"],
                        "dest_port": order["dest_port"],
                        "max_days": order.get("max_days"), "priority": order.get("priority")
                    }}],
                    "tool_results": [{"success": True, "tool": "compare_freight",
                                    "total_plans": 0}],
                    "model": active_llm.model, "configured": True,
                    "session_id": new_session_id, "parse_source": parse_source,
                    "feedback_source": fb["feedback_source"],
                    "feedback_reason": fb["feedback_reason"],
                }

        # ---- 3c: get_ports / get_statistics → 直接调工具 ----
        if intent in ("get_ports", "get_statistics"):
            _clear_session()
            tool_name = intent
            try:
                tool_result = await tool_manager.execute_tool(tool_name, {})
            except Exception:
                tool_result = {"success": False, "tool": tool_name, "error": "工具执行失败"}

            if tool_result.get("success"):
                if intent == "get_ports":
                    data = tool_result["result"]
                    message = (
                        f"可用港口信息：\n\n"
                        f"起运港（{data.get('total_orig_ports', 0)}个）：\n"
                        f"{', '.join(data.get('orig_ports', []))}\n\n"
                        f"目的港（{data.get('total_dest_ports', 0)}个）：\n"
                        f"{', '.join(data.get('dest_ports', []))}\n\n"
                        f"您可以输入'从[起运港]运[重量]到[目的港]'来查询运费方案。"
                    )
                else:
                    data = tool_result["result"]
                    message = (
                        f"系统数据统计：\n\n"
                        f"总报价记录：{data.get('total_records', 0)} 条\n"
                        f"承运商数量：{data.get('total_carriers', 0)} 家\n"
                        f"运输方式：{', '.join(data.get('transport_modes', []))}\n"
                        f"服务级别：{', '.join(data.get('service_levels', []))}"
                    )

                return {
                    "reply_type": "general",
                    "message": message,
                    "intent": intent,
                    "missing_fields": [], "order": None,
                    "recommendation": None, "plans": [],
                    "next_actions": ["输入运输需求查询运费方案"],
                    "response": message,
                    "tool_calls": [{"tool": tool_name, "parameters": {}}],
                    "tool_results": [tool_result],
                    "model": active_llm.model, "configured": True,
                    "session_id": new_session_id, "parse_source": parse_source
                }
            else:
                return {
                    "reply_type": "error",
                    "message": f"查询失败：{tool_result.get('error', '未知错误')}",
                    "intent": intent,
                    "missing_fields": [], "order": None,
                    "recommendation": None, "plans": [],
                    "next_actions": ["重试查询"],
                    "response": "查询失败",
                    "tool_calls": [], "tool_results": [tool_result],
                    "model": active_llm.model, "configured": True,
                    "session_id": new_session_id, "parse_source": parse_source
                }

        # ---- 3d: explain_cost / export_report / compare_carriers → LLM工具调用 ----
        if intent in ("explain_cost", "export_report", "compare_carriers"):
            _clear_session()
            llm_result = active_llm.chat_with_tools(request.message)
            tool_calls = llm_result.get("tool_calls", [])
            tool_results = []
            if tool_calls and tool_manager:
                for tc in tool_calls:
                    try:
                        tr = await tool_manager.execute_tool(tc.get("tool"), tc.get("parameters", {}))
                        tool_results.append(tr)
                    except Exception as e:
                        tool_results.append({"success": False, "tool": tc.get("tool"), "error": str(e)})

            return {
                "reply_type": "general",
                "message": llm_result.get("response", ""),
                "intent": intent,
                "missing_fields": missing,
                "order": order_info.model_dump() if order_info else None,
                "recommendation": None, "plans": [],
                "next_actions": [],
                "response": llm_result.get("response", ""),
                "tool_calls": tool_calls, "tool_results": tool_results,
                "model": active_llm.model, "configured": True,
                "session_id": llm_result.get("session_id"), "parse_source": parse_source
            }

        # ---- 3e: general 闲聊 → LLM 直接对话 ----
        _clear_session()
        llm_result = active_llm.chat_with_tools(request.message)
        tool_calls = llm_result.get("tool_calls", [])
        tool_results = []
        if tool_calls and tool_manager:
            for tc in tool_calls:
                try:
                    tr = await tool_manager.execute_tool(tc.get("tool"), tc.get("parameters", {}))
                    tool_results.append(tr)
                except Exception as e:
                    tool_results.append({"success": False, "tool": tc.get("tool"), "error": str(e)})

        return {
            "reply_type": "general",
            "message": llm_result.get("response", ""),
            "intent": "general",
            "missing_fields": [], "order": None,
            "recommendation": None, "plans": [],
            "next_actions": ["输入运输需求开始查询，例如：从大连运100kg到厦门"],
            "response": llm_result.get("response", ""),
            "tool_calls": tool_calls, "tool_results": tool_results,
            "model": active_llm.model, "configured": True,
            "session_id": llm_result.get("session_id"), "parse_source": parse_source
        }

    except Exception as e:
        return {
            "reply_type": "error",
            "message": f"处理请求时发生错误：{str(e)}",
            "intent": "general",
            "missing_fields": [], "order": None,
            "recommendation": None, "plans": [],
            "next_actions": ["重试请求", "检查输入格式"],
            "response": f"处理失败: {str(e)}",
            "tool_calls": [], "tool_results": [], "error": str(e),
            "model": active_llm.model if active_llm else "none",
            "configured": active_llm.is_configured() if active_llm else False,
            "session_id": None, "parse_source": "fallback"
        }


@app.get("/api/tools")
async def get_tools():
    """获取所有可用工具列表"""
    if tool_manager is None:
        return {"tools": [], "count": 0}
    return {
        "tools": tool_manager.get_tools_schema(),
        "count": len(tool_manager.tools)
    }


@app.post("/api/execute_tool")
async def execute_tool(tool_name: str, parameters: dict = {}):
    """直接执行指定工具"""
    if tool_manager is None:
        raise HTTPException(status_code=500, detail="工具管理器未初始化")
    try:
        result = await tool_manager.execute_tool(tool_name, parameters)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/parse")
async def parse_order(request: ParseRequest):
    """将自然语言描述解析为结构化订单数据，支持CoT思维链"""
    if llm_service is None:
        return {"error": "LLM服务未配置", "data": None}
    try:
        result = llm_service.parse_order(request.text, request.session_id)
        return {"error": None, "data": result}
    except Exception as e:
        return {"error": str(e), "data": None}


@app.post("/api/continue")
async def continue_conversation(request: ContinueRequest):
    """继续多轮对话，补充缺失信息"""
    if llm_service is None:
        return {"error": "LLM服务未配置", "data": None}
    try:
        result = llm_service.continue_conversation(request.session_id, request.message)
        return {"error": None, "data": result}
    except Exception as e:
        return {"error": str(e), "data": None}


@app.get("/api/session/{session_id}")
async def get_session_status(session_id: str):
    """获取会话状态"""
    if llm_service is None:
        return {"error": "LLM服务未配置", "data": None}
    try:
        result = llm_service.get_session_status(session_id)
        return {"error": None, "data": result}
    except Exception as e:
        return {"error": str(e), "data": None}


@app.post("/api/export")
async def export_report(order: DocxExportRequest):
    """导出比价报告"""
    try:
        result = freight_service.compare(order)
        report = generate_report(result, feedback=order.feedback, use_ai=order.use_ai)
        return {"report": report}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/status")
async def get_status(req: Request):
    """获取系统状态（支持检查用户自定义凭证）"""
    # 检查用户自定义凭证
    user_key = req.headers.get("X-API-Key")
    user_url = req.headers.get("X-Base-URL")
    user_model = req.headers.get("X-Model")

    if user_key and user_url:
        # 用户有自定义凭证，检查其有效性
        from llm_service import LLMService
        user_llm = LLMService(
            tool_manager=None,
            api_key=user_key,
            base_url=user_url,
            model=user_model or "mimo-v2.5-pro",
        )
        return {
            "data_source": active_data_source,
            "llm_configured": user_llm.is_configured(),
            "llm_model": user_llm.model,
            "has_custom_agent": True,
        }

    # 使用全局配置
    return {
        "data_source": active_data_source,
        "llm_configured": llm_service.is_configured() if llm_service else False,
        "llm_model": llm_service.model if llm_service else "none",
        "has_custom_agent": False,
    }


# ================================================================
# 管理接口
# ================================================================

@app.get("/admin/datasource")
async def get_datasource_info():
    """获取当前数据源信息"""
    info = {
        "active_source": active_data_source,
        "config_data_store": config.DATA_STORE,
        "db_path": config.DB_PATH,
    }
    if active_data_source == "sqlite":
        from database import check_db_integrity
        info["db_healthy"] = check_db_integrity()
    return info


@app.post("/admin/reload")
async def reload_datasource():
    """重新加载数据库（重新导入 CSV 并清缓存）"""
    global data_store, freight_service, active_data_source

    try:
        if config.DATA_STORE == "sqlite":
            from database import get_engine, get_session_factory, init_database, check_db_integrity
            from freight_service import DBDataStore, FreightService
            from db_models import FreightRate

            # 清空表并重新导入
            engine = get_engine()
            init_database()

            session = get_session_factory()()
            try:
                session.query(FreightRate).delete()
                session.commit()
            finally:
                session.close()

            from init_db import import_csv_to_db
            import_csv_to_db()

            if not check_db_integrity(config.EXPECTED_MIN_ROW_COUNT):
                raise RuntimeError("数据库完整性检查失败")

            session_factory = get_session_factory()
            data_store = DBDataStore(session_factory, auto_init=False)
            data_store._df = None  # 清除 DataFrame 缓存
            freight_service = FreightService(data_store)
            active_data_source = "sqlite"
        else:
            # CSV 模式: 重新加载
            _init_data_source()

        # 清除缓存
        if hasattr(data_store, 'clear_cache'):
            data_store.clear_cache()
        freight_service.clear_compare_cache()

        return {"success": True, "data_source": active_data_source, "message": "数据源已重新加载"}
    except Exception as e:
        logger.error(f"重新加载失败: {e}")
        return {"success": False, "data_source": active_data_source, "message": f"重新加载失败: {e}"}


def generate_report(result: ComparisonResult, feedback: str = None, use_ai: bool = True) -> str:
    """生成比价报告文本"""
    lines = []
    lines.append("=" * 60)
    lines.append("运输方案比价报告")
    lines.append("=" * 60)

    # 标注是否使用了AI助手
    if not use_ai:
        lines.append("")
        lines.append("【未使用AI助手】")
        lines.append("本报告由本地解析模式生成，未使用AI智能分析。")

    lines.append("")
    lines.append("【订单信息】")
    lines.append(f"  起运港: {result.order_info.orig_port}")
    lines.append(f"  目的港: {result.order_info.dest_port}")
    lines.append(f"  货物重量: {result.order_info.weight} kg")
    if result.order_info.max_days:
        lines.append(f"  时效要求: ≤{result.order_info.max_days}天")
    lines.append("")
    lines.append(f"【查询结果】共找到 {result.total_plans_found} 个可用方案")
    lines.append("")

    if result.available_plans:
        lines.append("【方案列表】")
        lines.append(f"{'承运商':<12} {'运输方式':<8} {'服务级别':<8} {'天数':<6} {'成本':<12} {'计算公式'}")
        lines.append("-" * 70)
        for plan in result.available_plans:
            mode_cn = "空运" if plan.mode == "AIR" else "陆运"
            service_cn = "门到门" if plan.service_level == "DTD" else "门到港"
            lines.append(f"{plan.carrier:<12} {mode_cn:<8} {service_cn:<8} {plan.transport_days:<6} ${plan.total_cost:<11.2f} {plan.cost_formula}")

    lines.append("")
    # 只有使用了AI助手且有feedback时才显示AI推荐理由
    if feedback and use_ai:
        lines.append("【AI 推荐理由】")
        lines.append(feedback)
        lines.append("")
    if result.recommended_plan:
        lines.append("【推荐方案】")
        lines.append(f"  承运商: {result.recommended_plan.plan.carrier}")
        mode_cn = "空运" if result.recommended_plan.plan.mode == "AIR" else "陆运"
        service_cn = "门到门" if result.recommended_plan.plan.service_level == "DTD" else "门到港"
        lines.append(f"  运输方式: {mode_cn}")
        lines.append(f"  服务级别: {service_cn}")
        lines.append(f"  运输天数: {result.recommended_plan.plan.transport_days}天")
        lines.append(f"  总成本: ${result.recommended_plan.plan.total_cost:.2f}")
        lines.append(f"  推荐理由: {result.recommended_plan.reason}")
    elif result.transfer_routes:
        lines.append("【转运方案】(未找到直达路线)")
        for i, tr in enumerate(result.transfer_routes):
            lines.append(f"  转运{i+1}: {' → '.join(tr.path)}")
            lines.append(f"    总成本: ${tr.total_cost:.2f} | 总耗时: {tr.total_estimated_days}天 | 经{tr.hop_count}次中转")
            for j, leg in enumerate(tr.legs):
                mode_cn = "空运" if leg.mode == "AIR" else "陆运"
                svc_cn = "门到门" if leg.service_level == "DTD" else "门到港"
                lines.append(f"    第{j+1}段 {leg.from_port}→{leg.to_port}: {leg.carrier} {mode_cn}({svc_cn}) ${leg.total_cost:.2f} / {leg.transport_days}天")
    elif result.fallback_transfer:
        lines.append("【次优推荐】")
        fb = result.fallback_transfer
        lines.append(f"  路径: {' → '.join(fb.path)}")
        lines.append(f"  总成本: ${fb.total_cost:.2f} | 总耗时: {fb.total_estimated_days}天")
        lines.append(f"  说明: {result.fallback_reason}")
    else:
        lines.append("【推荐方案】无满足条件的方案")

    lines.append("")
    lines.append("=" * 60)
    return "\n".join(lines)


# ================================================================
# Word 文档导出
# ================================================================

def generate_docx(result: ComparisonResult, feedback: str = None, use_ai: bool = True) -> bytes:
    """生成比价报告 Word 文档，返回 bytes"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from io import BytesIO

    doc = Document()

    # 标题
    title = doc.add_heading('运输方案比价报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 标注是否使用了AI助手
    if not use_ai:
        ai_notice = doc.add_paragraph()
        ai_notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ai_notice.add_run('【未使用AI助手】本报告由本地解析模式生成，未使用AI智能分析。')
        run.font.color.rgb = RGBColor(0xc2, 0x41, 0x0c)
        run.font.size = Pt(10)

    # 订单信息
    doc.add_heading('订单信息', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = table.rows[0].cells
    cells[0].text = '项目'
    cells[1].text = '内容'

    order_rows = [
        ('起运港', result.order_info.orig_port),
        ('目的港', result.order_info.dest_port),
        ('货物重量', f'{result.order_info.weight} kg'),
    ]
    if result.order_info.max_days:
        order_rows.append(('时效要求', f'≤ {result.order_info.max_days} 天'))
    if result.order_info.priority:
        pri_cn = '时效优先' if result.order_info.priority == 'time' else '成本优先'
        order_rows.append(('优先级', pri_cn))

    for label, value in order_rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value)

    # 方案列表
    doc.add_heading(f'查询结果（共 {result.total_plans_found} 个方案）', level=1)

    if result.available_plans:
        headers = ['承运商', '运输方式', '服务级别', '运输天数', '总成本', '综合评分']
        plan_table = doc.add_table(rows=1, cols=len(headers))
        plan_table.style = 'Light Grid Accent 1'
        plan_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            plan_table.rows[0].cells[i].text = h

        for plan in result.available_plans:
            row = plan_table.add_row()
            row.cells[0].text = plan.carrier
            row.cells[1].text = '空运' if plan.mode == 'AIR' else '陆运'
            row.cells[2].text = '门到门' if plan.service_level == 'DTD' else '门到港'
            row.cells[3].text = f'{plan.transport_days} 天'
            row.cells[4].text = f'${plan.total_cost:.2f}'
            row.cells[5].text = f'{plan.score:.3f}'

    # AI 反馈推荐（LLM 生成的推荐理由）- 只有使用了AI且有feedback时才显示
    if feedback and use_ai:
        doc.add_heading('AI 推荐理由', level=1)
        fb_para = doc.add_paragraph(feedback)
        fb_para.paragraph_format.space_after = Pt(12)

    # 推荐方案 / 转运方案
    if result.recommended_plan:
        doc.add_heading('推荐方案', level=1)
        rp = result.recommended_plan.plan
        mode_cn = '空运' if rp.mode == 'AIR' else '陆运'
        service_cn = '门到门' if rp.service_level == 'DTD' else '门到港'

        rec_table = doc.add_table(rows=6, cols=2)
        rec_table.style = 'Light Grid Accent 1'
        rec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        fields = [
            ('承运商', rp.carrier),
            ('运输方式', f'{mode_cn}（{service_cn}）'),
            ('运输天数', f'{rp.transport_days} 天'),
            ('总成本', f'${rp.total_cost:.2f}'),
            ('综合评分', f'{rp.score:.3f} / 1.0'),
            ('推荐理由', result.recommended_plan.reason),
        ]
        for i, (label, value) in enumerate(fields):
            rec_table.rows[i].cells[0].text = label
            rec_table.rows[i].cells[1].text = str(value)

    elif result.transfer_routes:
        doc.add_heading('转运方案（未找到直达路线）', level=1)
        for i, tr in enumerate(result.transfer_routes):
            doc.add_heading(f'转运路线 {i+1}: {" → ".join(tr.path)}', level=2)
            tr_meta = doc.add_table(rows=3, cols=2)
            tr_meta.style = 'Light Grid Accent 1'
            tr_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
            meta_fields = [
                ('总成本', f'${tr.total_cost:.2f}'),
                ('总耗时', f'{tr.total_estimated_days} 天（运输 {tr.total_days} 天 + 转运 {tr.hop_count} 天）'),
                ('综合评分', f'{tr.score:.3f} / 1.0' if tr.score is not None else '未评分'),
            ]
            for j, (label, value) in enumerate(meta_fields):
                tr_meta.rows[j].cells[0].text = label
                tr_meta.rows[j].cells[1].text = str(value)

            # 每段明细
            doc.add_paragraph('')
            leg_headers = ['段次', '路线', '承运商', '运输方式', '服务级别', '运输天数', '费用']
            leg_table = doc.add_table(rows=1, cols=len(leg_headers))
            leg_table.style = 'Light Grid Accent 1'
            leg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for k, h in enumerate(leg_headers):
                leg_table.rows[0].cells[k].text = h

            for j, leg in enumerate(tr.legs):
                row = leg_table.add_row()
                row.cells[0].text = f'第{j+1}段'
                row.cells[1].text = f'{leg.from_port} → {leg.to_port}'
                row.cells[2].text = leg.carrier
                row.cells[3].text = '空运' if leg.mode == 'AIR' else '陆运'
                row.cells[4].text = '门到门' if leg.service_level == 'DTD' else '门到港'
                row.cells[5].text = f'{leg.transport_days} 天'
                row.cells[6].text = f'${leg.total_cost:.2f}'

    elif result.fallback_transfer:
        doc.add_heading('次优推荐', level=1)
        fb = result.fallback_transfer
        fb_p = doc.add_paragraph(result.fallback_reason or '当前条件下无可用方案，以下为最接近的转运路线')
        fb_p.paragraph_format.space_after = Pt(12)

        fb_table = doc.add_table(rows=3, cols=2)
        fb_table.style = 'Light Grid Accent 1'
        fb_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        fb_fields = [
            ('路径', ' → '.join(fb.path)),
            ('总成本', f'${fb.total_cost:.2f}'),
            ('总耗时', f'{fb.total_estimated_days} 天'),
        ]
        for j, (label, value) in enumerate(fb_fields):
            fb_table.rows[j].cells[0].text = label
            fb_table.rows[j].cells[1].text = str(value)
    else:
        doc.add_heading('推荐方案', level=1)
        doc.add_paragraph('无满足条件的方案')

    # 页脚
    doc.add_paragraph('')
    footer = doc.add_paragraph('本报告由运输方案比价与优化智能体自动生成')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    footer.runs[0].font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_history_docx(items: list) -> bytes:
    """生成历史查询记录 Word 文档，返回 bytes"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from io import BytesIO

    doc = Document()

    title = doc.add_heading('历史查询记录', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary = doc.add_paragraph(f'共 {len(items)} 条记录')
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    for idx, item in enumerate(items, 1):
        # 序号标题
        doc.add_heading(f'记录 {idx}', level=2)

        # 基本信息表
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_rows = [
            ('时间', item.get('timestamp', '')),
            ('用户输入', item.get('userInput', '')),
            ('结果类型', item.get('replyType', '')),
        ]
        for i, (label, value) in enumerate(info_rows):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = str(value)

        # 订单信息（如果有）
        order = item.get('order')
        if order:
            order_table = doc.add_table(rows=0, cols=2)
            order_table.style = 'Light Grid Accent 1'
            order_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            order_fields = [
                ('起运港', order.get('orig_port', '')),
                ('目的港', order.get('dest_port', '')),
                ('重量', f"{order.get('weight', '')} kg" if order.get('weight') else ''),
                ('时效', f"{order.get('max_days', '')} 天" if order.get('max_days') else ''),
            ]
            for label, value in order_fields:
                if value:
                    row = order_table.add_row()
                    row.cells[0].text = label
                    row.cells[1].text = str(value)

        # 结果
        result_str = item.get('result', '')
        message = item.get('message', '')
        if result_str:
            p = doc.add_paragraph()
            p.add_run('查询结果：').bold = True
            p.add_run(result_str)
        elif message:
            p = doc.add_paragraph()
            p.add_run('系统回复：').bold = True
            p.add_run(message[:200])

        # 分隔
        if idx < len(items):
            doc.add_paragraph('')

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


@app.post("/api/export_docx")
async def export_docx(order: DocxExportRequest):
    """导出比价报告 Word 文档"""
    from fastapi.responses import Response
    from urllib.parse import quote
    try:
        result = freight_service.compare(order)
        docx_bytes = generate_docx(result, feedback=order.feedback, use_ai=order.use_ai)
        filename = f"比价报告_{order.orig_port}_{order.dest_port}_{order.weight}kg.docx"
        encoded = quote(filename)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class HistoryExportRequest(BaseModel):
    items: List[dict]


@app.post("/api/export_history_docx")
async def export_history_docx(request: HistoryExportRequest):
    """导出历史查询记录 Word 文档"""
    from fastapi.responses import Response
    from urllib.parse import quote
    try:
        docx_bytes = generate_history_docx(request.items)
        encoded = quote('历史查询记录.docx')
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
